"""
다양한 문서 및 첨부 파일 파싱 모듈 (PDF, DOCX, HWPX, HWP, TXT, 이미지 등)
"""

import os
import zipfile
import xml.etree.ElementTree as ET

def extract_text_from_file(file_path: str) -> dict:
    """
    파일 경로를 받아 텍스트 및 메타데이터를 추출하여 dict로 반환
    반환 형태: {"type": "text"|"image", "content": str|bytes, "filename": str, "extension": str}
    """
    if not os.path.exists(file_path):
        raise FileNotFoundError(f"파일을 찾을 수 없습니다: {file_path}")

    filename = os.path.basename(file_path)
    ext = os.path.splitext(filename)[1].lower()

    # 1. PDF (.pdf)
    if ext == ".pdf":
        text = _read_pdf(file_path)
        return {"type": "text", "content": text, "filename": filename, "extension": ext}

    # 2. Word (.docx)
    elif ext == ".docx":
        text = _read_docx(file_path)
        return {"type": "text", "content": text, "filename": filename, "extension": ext}

    # 3. 한글 HWPX (.hwpx)
    elif ext == ".hwpx":
        text = _read_hwpx(file_path)
        return {"type": "text", "content": text, "filename": filename, "extension": ext}

    # 4. 한글 HWP (.hwp)
    elif ext == ".hwp":
        text = _read_hwp_text(file_path)
        return {"type": "text", "content": text, "filename": filename, "extension": ext}

    # 5. 이미지 계열 (.png, .jpg, .jpeg, .webp)
    elif ext in [".png", ".jpg", ".jpeg", ".webp"]:
        with open(file_path, "rb") as f:
            image_bytes = f.read()
        return {"type": "image", "content": image_bytes, "filename": filename, "extension": ext, "path": file_path}

    # 6. 일반 텍스트 (.txt, .md, .csv, .json) 및 기타 파일
    else:
        text = _read_plain_text(file_path)
        return {"type": "text", "content": text, "filename": filename, "extension": ext}


def _read_plain_text(file_path: str) -> str:
    """UTF-8 또는 CP949 인코딩으로 일반 텍스트 파일 읽기"""
    encodings = ["utf-8", "cp949", "euc-kr", "utf-16"]
    for enc in encodings:
        try:
            with open(file_path, "r", encoding=enc) as f:
                return f.read()
        except UnicodeDecodeError:
            continue
    with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
        return f.read()


def _read_pdf(file_path: str) -> str:
    """PDF 파일의 텍스트 추출"""
    # pyrefly: ignore [missing-import]
    from pypdf import PdfReader
    reader = PdfReader(file_path)
    text_list = []
    for idx, page in enumerate(reader.pages):
        page_text = page.extract_text()
        if page_text:
            text_list.append(f"--- [페이지 {idx+1}] ---\n{page_text.strip()}")
    return "\n\n".join(text_list)


def _read_docx(file_path: str) -> str:
    """Word DOCX 파일의 텍스트 추출"""
    # pyrefly: ignore [missing-import]
    import docx
    doc = docx.Document(file_path)
    text_list = []
    for p in doc.paragraphs:
        if p.text.strip():
            text_list.append(p.text.strip())
    # 표 내용도 추출
    for table in doc.tables:
        for row in table.rows:
            row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if row_text:
                text_list.append(" | ".join(row_text))
    return "\n".join(text_list)


def _read_hwpx(file_path: str) -> str:
    """HWPX (ZIP 기반 XML) 파일의 텍스트 추출"""
    text_list = []
    try:
        with zipfile.ZipFile(file_path, 'r') as z:
            for name in z.namelist():
                if name.startswith("Contents/section") and name.endswith(".xml"):
                    xml_content = z.read(name)
                    tree = ET.fromstring(xml_content)
                    # XML의 텍스트 노드 추출
                    for elem in tree.iter():
                        if elem.tag.endswith("t") and elem.text:
                            text_list.append(elem.text)
    except Exception as e:
        return f"[HWPX 읽기 오류: {e}]"
    return "\n".join(text_list)


def _read_hwp_text(file_path: str) -> str:
    """HWP 5.0 바이너리 파일 텍스트 추출 (가능한 텍스트 파싱)"""
    try:
        import olefile
        if olefile.isOleFile(file_path):
            ole = olefile.OleFileIO(file_path)
            # BodyText 섹션 검색
            dirs = ole.listdir()
            text_list = []
            for d in dirs:
                if d[0] == "BodyText":
                    stream = ole.openstream(d)
                    data = stream.read()
                    try:
                        import zlib
                        decompressed = zlib.decompress(data, -15)
                        text_list.append(decompressed.decode("utf-16le", errors="ignore"))
                    except Exception:
                        pass
            if text_list:
                return "\n".join(text_list)
    except Exception:
        pass
    
    # 대체 방식: 일반 바이너리에서 문자열 추출
    try:
        with open(file_path, "rb") as f:
            raw = f.read()
        return raw.decode("utf-16le", errors="ignore")
    except Exception as e:
        return f"[HWP 파싱 제한: {e}]"
